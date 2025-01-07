import streamlit as st
import sqlite3
import pandas as pd
from io import BytesIO
from docx import Document
import xlsxwriter
from fpdf import FPDF
conn = sqlite3.connect("employee_data.db")
c = conn.cursor()

c.execute('''CREATE TABLE IF NOT EXISTS employees (
                id INTEGER PRIMARY KEY,
                name TEXT,
                position TEXT,
                department TEXT,
                hire_date TEXT,
                performance_review TEXT
            )''')
conn.commit()

def add_employee(name, position, department, hire_date, review):
    c.execute("INSERT INTO employees (name, position, department, hire_date, performance_review) VALUES (?, ?, ?, ?, ?)",
              (name, position, department, hire_date, review))
    conn.commit()

def get_employees():
    c.execute("SELECT * FROM employees")
    return c.fetchall()

def update_employee(emp_id, name, position, department, hire_date, review):
    c.execute("UPDATE employees SET name=?, position=?, department=?, hire_date=?, performance_review=? WHERE id=?",
              (name, position, department, hire_date, review, emp_id))
    conn.commit()

def delete_employee(emp_id):
    c.execute("DELETE FROM employees WHERE id=?", (emp_id,))
    conn.commit()

def export_to_excel(data):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df = pd.DataFrame(data, columns=['ID', 'Name', 'Position', 'Department', 'Hire Date', 'Performance Review'])
        df.to_excel(writer, index=False, sheet_name='Employees')
    output.seek(0)
    return output

def export_to_word(data):
    doc = Document()
    doc.add_heading('Employee Data', level=1)
    for row in data:
        doc.add_paragraph(f"ID: {row[0]}\nName: {row[1]}\nPosition: {row[2]}\nDepartment: {row[3]}\nHire Date: {row[4]}\nPerformance Review: {row[5]}\n")
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def export_to_pdf(data):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Employee Data", ln=True, align='C')
    for row in data:
        pdf.cell(200, 10, txt=f"ID: {row[0]} | Name: {row[1]} | Position: {row[2]} | Department: {row[3]} | Hire Date: {row[4]} | Performance Review: {row[5]}", ln=True)
    output = BytesIO()
    pdf.output(output)
    output.seek(0)
    return output
st.sidebar.title("Menu")
menu = st.sidebar.radio("Choose an option:", ["Add Employee", "View Employees", "Export Data", "Search Employees", "Generate Reports"])

if menu == "Add Employee":
    st.title("Add New Employee")
    with st.form("add_employee_form"):
        name = st.text_input("Name")
        position = st.text_input("Position")
        department = st.text_input("Department")
        hire_date = st.date_input("Hire Date")
        review = st.text_area("Performance Review")
        submitted = st.form_submit_button("Add Employee")

        if submitted:
            add_employee(name, position, department, str(hire_date), review)
            st.success("Employee added successfully!")

elif menu == "View Employees":
    st.title("Employee Records")
    employees = get_employees()
    df = pd.DataFrame(employees, columns=['ID', 'Name', 'Position', 'Department', 'Hire Date', 'Performance Review'])
    st.dataframe(df)

    if st.checkbox("Edit Employee"):
        emp_id = st.number_input("Employee ID", min_value=1, step=1)
        with st.form("edit_employee_form"):
            name = st.text_input("Name")
            position = st.text_input("Position")
            department = st.text_input("Department")
            hire_date = st.date_input("Hire Date")
            review = st.text_area("Performance Review")
            submitted = st.form_submit_button("Update Employee")

            if submitted:
                update_employee(emp_id, name, position, department, str(hire_date), review)
                st.success("Employee updated successfully!")

    if st.checkbox("Delete Employee"):
        emp_id = st.number_input("Employee ID to delete", min_value=1, step=1)
        if st.button("Delete"):
            delete_employee(emp_id)
            st.success("Employee deleted successfully!")

elif menu == "Export Data":
    st.title("Export Employee Data")
    employees = get_employees()
    if st.button("Export to Excel"):
        excel_data = export_to_excel(employees)
        st.download_button(label="Download Excel File", data=excel_data, file_name="employee_data.xlsx")

    if st.button("Export to Word"):
        word_data = export_to_word(employees)
        st.download_button(label="Download Word File", data=word_data, file_name="employee_data.docx")

    if st.button("Export to PDF"):
        pdf_data = export_to_pdf(employees)
        st.download_button(label="Download PDF File", data=pdf_data, file_name="employee_data.pdf")

elif menu == "Search Employees":
    st.title("Search Employees")
    search_term = st.text_input("Enter search term")
    if search_term:
        c.execute("SELECT * FROM employees WHERE name LIKE ? OR position LIKE ? OR department LIKE ?", (f"%{search_term}%", f"%{search_term}%", f"%{search_term}%"))
        results = c.fetchall()
        if results:
            st.write(pd.DataFrame(results, columns=['ID', 'Name', 'Position', 'Department', 'Hire Date', 'Performance Review']))
        else:
            st.write("No matching employees found.")

elif menu == "Generate Reports":
    st.title("Generate Reports")
    report_type = st.selectbox("Select report type", ["Summary Report", "Detailed Report"])
    employees = get_employees()

    if report_type == "Summary Report":
        st.write("Total Employees: ", len(employees))
        departments = pd.DataFrame(employees, columns=['ID', 'Name', 'Position', 'Department', 'Hire Date', 'Performance Review'])['Department'].value_counts()
        st.bar_chart(departments)

    elif report_type == "Detailed Report":
        st.write(pd.DataFrame(employees, columns=['ID', 'Name', 'Position', 'Department', 'Hire Date', 'Performance Review']))
conn.close()
